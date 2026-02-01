import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

# Читаем Excel-файл
df = pd.read_excel('480.xlsx', header=None, engine='openpyxl')

#print(df)
#print(df.info())
#print(df.describe())

#print(df.loc[100:112])
#print(df.loc[0])
print('Количество строк:', df[0].count())

# Открываем Word-документ
doc = Document('table20sheets.docx')

tables = doc.tables
print('Таблиц в файле:', len(tables)) # должно быть 40 для 480 слов

for i in range(df[0].count()):
        #заполняем англ. слово
    #вычислим номер таблицы ряда и строки
    table = i // 24
    row = i // 3 - table * 8
    column = i  - table * 24 - row * 3
    #print('i:', i, 
    #      'table', table,
    #      'row', row,
    #      'column', column,
    #      df.loc[i][0])
    
    cell = tables[table].cell(row, column)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(df.loc[i][0])
    run.font.size = Pt(14)
    
        #заполняем обратную сторону
    #вычислим номер таблицы ряда и строки
    table = len(tables) - table - 1
    row = 7 - row
    #column = column
    
    #транскрипция
    cell = tables[table].cell(row, column)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6) # Отступ  снизу (между абзацами)
    run = paragraph.add_run(df.loc[i][2])
    run.font.size = Pt(12)
    
    #перевод
    cell = tables[table].cell(row, column)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)  # вместо None
    paragraph.paragraph_format.left_indent = Cm(0.5) # Отступ слева и справа
    paragraph.paragraph_format.right_indent = Cm(0.5) # Отступ слева и справа
    paragraph.paragraph_format.space_after = Pt(6) # Отступ  снизу (между абзацами)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # говорим, что используем коэффициент
    paragraph.paragraph_format.line_spacing = 1
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(df.loc[i][1])
    run.font.size = Pt(12)
    
    #пример
    cell = tables[table].cell(row, column)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)  # вместо None
    paragraph.paragraph_format.left_indent = Cm(0.5) # Отступ слева и справа
    paragraph.paragraph_format.right_indent = Cm(0.5) # Отступ слева и справа
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # говорим, что используем коэффициент
    paragraph.paragraph_format.line_spacing = 1
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(df.loc[i][3])
    run.font.size = Pt(12)
    
# Сохраняем результат
doc.save('table20sheets_filled.docx')